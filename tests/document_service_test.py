"""Local tests for PDF/DOCX validation, extraction, and retrieval."""

import io
import tempfile
from pathlib import Path
from unittest import TestCase
from unittest.mock import patch

from docx import Document
from fastapi.testclient import TestClient

from file_service.main import app


class DocumentServiceTest(TestCase):
    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory(dir=r"C:\tmp")
        self.storage_patch = patch(
            "file_service.storage.FILE_STORAGE_DIR",
            Path(self.temp_dir.name),
        )
        self.storage_patch.start()
        self.client = TestClient(app)

    def tearDown(self):
        self.client.close()
        self.storage_patch.stop()
        self.temp_dir.cleanup()

    @staticmethod
    def docx_bytes():
        buffer = io.BytesIO()
        document = Document()
        document.add_heading("Attention test", level=1)
        document.add_paragraph(
            "Scaled dot-product attention divides logits by the square root "
            "of the key dimension."
        )
        document.save(buffer)
        return buffer.getvalue()

    def test_docx_upload_extract_and_retrieve(self):
        response = self.client.post(
            "/documents",
            files={
                "file": (
                    "attention.docx",
                    self.docx_bytes(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        self.assertEqual(response.status_code, 201, response.text)
        uploaded = response.json()
        self.assertIn("Scaled dot-product attention", uploaded["extracted_text"])
        self.assertGreater(uploaded["character_count"], 20)

        fetched = self.client.get(f"/documents/{uploaded['id']}")
        self.assertEqual(fetched.status_code, 200)
        self.assertEqual(
            fetched.json()["extracted_text"], uploaded["extracted_text"]
        )

        source = self.client.get(f"/documents/{uploaded['id']}/file")
        self.assertEqual(source.status_code, 200)
        self.assertTrue(source.content.startswith(b"PK"))
        self.assertIn(
            "attachment", source.headers.get("content-disposition", "")
        )

    def test_rejects_non_document_upload(self):
        response = self.client.post(
            "/documents",
            files={"file": ("notes.txt", b"not allowed", "text/plain")},
        )
        self.assertEqual(response.status_code, 415)


if __name__ == "__main__":
    import unittest

    unittest.main()
