import json
import re
import shutil
from datetime import datetime, timezone
from pathlib import Path
from uuid import UUID, uuid4

from docx import Document as WordDocument
from pypdf import PdfReader

from file_service.config import FILE_STORAGE_DIR

ALLOWED_EXTENSIONS = {".pdf": "pdf", ".docx": "docx"}


class DocumentStorageError(Exception):
    pass


def initialize_storage():
    FILE_STORAGE_DIR.mkdir(parents=True, exist_ok=True)


def _safe_filename(filename):
    basename = Path(filename or "document").name
    cleaned = re.sub(r"[^A-Za-z0-9._ -]", "_", basename).strip(" .")
    return cleaned or "document"


def validate_filename(filename):
    safe_name = _safe_filename(filename)
    extension = Path(safe_name).suffix.lower()
    if extension not in ALLOWED_EXTENSIONS:
        raise DocumentStorageError("Only PDF and DOCX files are supported")
    return safe_name, extension, ALLOWED_EXTENSIONS[extension]


def _extract_pdf(path):
    try:
        reader = PdfReader(str(path))
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception as exc:
                raise DocumentStorageError(
                    "Password-protected PDFs are not supported"
                ) from exc
        pages = []
        for page in reader.pages:
            text = page.extract_text() or ""
            if text.strip():
                pages.append(text.strip())
        return "\n\n".join(pages)
    except DocumentStorageError:
        raise
    except Exception as exc:
        raise DocumentStorageError("Unable to read this PDF") from exc


def _extract_docx(path):
    try:
        document = WordDocument(str(path))
        blocks = [
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    blocks.append("\t".join(cells))
        return "\n\n".join(blocks)
    except Exception as exc:
        raise DocumentStorageError("Unable to read this DOCX file") from exc


def _extract_text(path, file_type):
    text = _extract_pdf(path) if file_type == "pdf" else _extract_docx(path)
    text = text.replace("\x00", "").strip()
    if not text:
        raise DocumentStorageError(
            "The document contains no extractable text; scanned PDFs need OCR"
        )
    return text


def save_document(filename, content_type, content):
    safe_name, extension, file_type = validate_filename(filename)
    document_id = uuid4()
    document_dir = FILE_STORAGE_DIR / str(document_id)
    source_path = document_dir / f"source{extension}"
    text_path = document_dir / "extracted.txt"
    metadata_path = document_dir / "metadata.json"

    document_dir.mkdir(parents=True, exist_ok=False)
    try:
        source_path.write_bytes(content)
        extracted_text = _extract_text(source_path, file_type)
        text_path.write_text(extracted_text, encoding="utf-8")
        metadata = {
            "id": str(document_id),
            "filename": safe_name,
            "file_type": file_type,
            "content_type": content_type or (
                "application/pdf"
                if file_type == "pdf"
                else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ),
            "size_bytes": len(content),
            "character_count": len(extracted_text),
            "created_at": datetime.now(timezone.utc).isoformat(),
        }
        metadata_path.write_text(
            json.dumps(metadata, indent=2), encoding="utf-8"
        )
        return {**metadata, "extracted_text": extracted_text}
    except Exception:
        shutil.rmtree(document_dir, ignore_errors=True)
        raise


def _document_dir(document_id):
    try:
        normalized = str(UUID(str(document_id)))
    except (TypeError, ValueError) as exc:
        raise FileNotFoundError("Document not found") from exc
    return FILE_STORAGE_DIR / normalized


def get_document(document_id):
    document_dir = _document_dir(document_id)
    metadata_path = document_dir / "metadata.json"
    text_path = document_dir / "extracted.txt"
    if not metadata_path.is_file() or not text_path.is_file():
        raise FileNotFoundError("Document not found")
    metadata = json.loads(metadata_path.read_text(encoding="utf-8"))
    return {
        **metadata,
        "extracted_text": text_path.read_text(encoding="utf-8"),
    }


def list_documents():
    initialize_storage()
    documents = []
    for metadata_path in FILE_STORAGE_DIR.glob("*/metadata.json"):
        try:
            documents.append(
                json.loads(metadata_path.read_text(encoding="utf-8"))
            )
        except (OSError, json.JSONDecodeError):
            continue
    return sorted(
        documents, key=lambda item: item.get("created_at", ""), reverse=True
    )


def delete_document(document_id):
    document_dir = _document_dir(document_id)
    if not document_dir.is_dir():
        return False
    shutil.rmtree(document_dir)
    return True
